"""
Serviço de geração de DOCX para relatórios OPME.
Usa python-docx para gerar documento Word profissional.
"""
import io
import uuid
import logging
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# ── Color palette ──
NAVY = RGBColor(0x1A, 0x3C, 0x6E)
DARK = RGBColor(0x2D, 0x2D, 0x2D)
BODY = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x6B, 0x6B, 0x6B)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
GREEN = RGBColor(0x27, 0xAE, 0x60)
RED = RGBColor(0xC0, 0x39, 0x2B)

FONT = "Calibri"


def _set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def _set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right with values like ('single', '4', '1A3C6E')."""
    tc_pr = cell._element.get_or_add_tcPr()
    tc_borders = tc_pr.makeelement(qn("w:tcBorders"), {})
    for edge, (style, sz, color) in kwargs.items():
        edge_el = tc_borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): style,
            qn("w:sz"): sz,
            qn("w:color"): color,
            qn("w:space"): "0",
        })
        tc_borders.append(edge_el)
    tc_pr.append(tc_borders)


def _add_run(paragraph, text, bold=False, size=11, color=BODY, font_name=FONT):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font_name
    # Force Calibri for East Asian/Complex Script as well
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    return run


def _add_bottom_border(paragraph, color="1A3C6E", sz="8"):
    """Add a bottom border to a paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): sz,
        qn("w:color"): color,
        qn("w:space"): "1",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _split_paragraphs(text: str) -> list[str]:
    if not text:
        return []
    paragraphs = []
    for block in text.split("\n\n"):
        clean = block.strip()
        if clean:
            paragraphs.append(clean.replace("\n", " "))
    if not paragraphs and text.strip():
        paragraphs = [text.strip()]
    return paragraphs


def generate_docx_bytes(
    justificativa: str,
    paciente_nome: str,
    cid: str,
    diagnostico_resumo: str,
    produto_nome: str,
    convenio: str = "",
    especialidade: str = "",
    codigo_tuss: str = "",
    referencias: list[str] = None,
    checklist: dict = None,
    medico_nome: str = "",
    medico_crm: str = "",
    medico_rqe: str = "",
    aprovado: bool = True,
    falha_terapeutica: str = "",
    risco_nao_realizacao: str = "",
    base_legal: str = "",
    clinica_nome: str = "",
    clinica_logo_url: str = "",
    paciente_dob: str = "",
    paciente_carteirinha: str = "",
    paciente_cpf: str = "",
    guia_numero: str = "",
    atendimento_numero: str = "",
    cids_secundarios: list = None,
    materiais_tuss: list = None,
    registro_anvisa: str = "",
    compliance_texto: str = "",
    **_extra,
) -> bytes:
    # Sanitize inputs
    justificativa = justificativa or ""
    paciente_nome = paciente_nome or ""
    cid = cid or ""
    diagnostico_resumo = diagnostico_resumo or ""
    produto_nome = produto_nome or ""
    convenio = convenio or ""
    especialidade = especialidade or ""
    codigo_tuss = codigo_tuss or ""
    referencias = referencias or []
    checklist = checklist or {}
    medico_nome = medico_nome or ""
    medico_crm = medico_crm or ""
    medico_rqe = medico_rqe or ""
    falha_terapeutica = falha_terapeutica or ""
    risco_nao_realizacao = risco_nao_realizacao or ""
    base_legal = base_legal or ""
    clinica_nome = clinica_nome or ""
    cids_secundarios = cids_secundarios or []
    materiais_tuss = materiais_tuss or []

    doc = Document()

    # ── Global style ──
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    style.font.color.rgb = BODY
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    protocolo = f"OPME-{datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:6].upper()}"

    # ══════════════════════════════════════════════════════════════════
    # HEADER
    # ══════════════════════════════════════════════════════════════════
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run(p_title, (clinica_nome or "RELATÓRIO MÉDICO COMPLEMENTAR").upper(), bold=True, size=15, color=NAVY)
    p_title.paragraph_format.space_after = Pt(0)

    p_sub = doc.add_paragraph()
    _add_run(p_sub, "Justificativa Técnica para Autorização de Material OPME", size=10, color=GRAY)
    p_sub.paragraph_format.space_after = Pt(4)
    _add_bottom_border(p_sub)

    # Protocol & date line
    p_proto = doc.add_paragraph()
    p_proto.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(p_proto, f"Data: {datetime.now().strftime('%d/%m/%Y')}  •  Protocolo: {protocolo}", size=9, color=LIGHT_GRAY)
    p_proto.paragraph_format.space_after = Pt(10)

    # ══════════════════════════════════════════════════════════════════
    # METADATA TABLE
    # ══════════════════════════════════════════════════════════════════
    meta_items = [
        ("Paciente", paciente_nome or "Não informado"),
        ("Nascimento", paciente_dob or "—"),
        ("Convênio", convenio or "Não informado"),
        ("Carteirinha", paciente_carteirinha or "—"),
        ("Diagnóstico (CID)", f"{cid} — {diagnostico_resumo}" if cid else diagnostico_resumo or "—"),
        ("Guia / Atend.", (guia_numero or "—") + (f" / {atendimento_numero}" if atendimento_numero else "")),
        ("Especialidade", especialidade or "—"),
        ("Código TUSS", codigo_tuss or "—"),
        ("Material OPME", produto_nome or "—"),
        ("Reg. ANVISA", registro_anvisa or "—"),
    ]
    if cids_secundarios:
        meta_items.append(("CID secundários", ", ".join(str(c) for c in cids_secundarios)))
        meta_items.append(("", ""))

    n_rows = (len(meta_items) + 1) // 2
    table = doc.add_table(rows=n_rows, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(meta_items):
        row_idx = i // 2
        col_offset = (i % 2) * 2
        cell_label = table.cell(row_idx, col_offset)
        cell_value = table.cell(row_idx, col_offset + 1)

        _set_cell_shading(cell_label, "F0F4FA")
        _set_cell_shading(cell_value, "F0F4FA")

        p_l = cell_label.paragraphs[0]
        p_l.paragraph_format.space_before = Pt(3)
        p_l.paragraph_format.space_after = Pt(3)
        _add_run(p_l, label, bold=True, size=9, color=NAVY)

        p_v = cell_value.paragraphs[0]
        p_v.paragraph_format.space_before = Pt(3)
        p_v.paragraph_format.space_after = Pt(3)
        _add_run(p_v, value, size=9, color=DARK)

    # Remove table borders for cleaner look
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell,
                top=("single", "2", "D8E2EF"),
                bottom=("single", "2", "D8E2EF"),
                left=("none", "0", "FFFFFF"),
                right=("none", "0", "FFFFFF"),
            )

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ══════════════════════════════════════════════════════════════════
    # HELPER: Add a section with title + body paragraphs
    # ══════════════════════════════════════════════════════════════════
    def _add_section(title: str, text: str):
        if not text or not text.strip():
            return
        p_heading = doc.add_paragraph()
        p_heading.paragraph_format.space_before = Pt(14)
        p_heading.paragraph_format.space_after = Pt(4)
        _add_run(p_heading, title.upper(), bold=True, size=10, color=NAVY)
        _add_bottom_border(p_heading, color="D8E2EF", sz="4")

        for para_text in _split_paragraphs(text):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Cm(1.0)
            _add_run(p, para_text, size=11, color=BODY)

    # ══════════════════════════════════════════════════════════════════
    # BODY SECTIONS (parseadas do corpo com títulos próprios)
    # ══════════════════════════════════════════════════════════════════
    from app.services.pdf_generator import _split_into_sections
    for sec in _split_into_sections(justificativa):
        _add_section(sec["titulo"], "\n\n".join(sec["paragrafos"]))

    if compliance_texto:
        _add_section("Adequação ao Rol / DUT (ANS)", compliance_texto)

    if base_legal:
        _add_section("Fundamentação Legal", base_legal)

    # ══════════════════════════════════════════════════════════════════
    # REFERÊNCIAS BIBLIOGRÁFICAS
    # ══════════════════════════════════════════════════════════════════
    if referencias:
        p_ref_title = doc.add_paragraph()
        p_ref_title.paragraph_format.space_before = Pt(16)
        p_ref_title.paragraph_format.space_after = Pt(6)
        _add_run(p_ref_title, "REFERÊNCIAS BIBLIOGRÁFICAS", bold=True, size=10, color=NAVY)
        _add_bottom_border(p_ref_title, color="D8E2EF", sz="4")

        for i, ref in enumerate(referencias, 1):
            p_ref = doc.add_paragraph()
            p_ref.paragraph_format.space_after = Pt(2)
            p_ref.paragraph_format.left_indent = Cm(0.5)
            if isinstance(ref, dict):
                texto = ref.get("texto", str(ref))
            else:
                texto = str(ref)
            _add_run(p_ref, f"{i}. {texto}", size=9, color=GRAY)

    # ══════════════════════════════════════════════════════════════════
    # CHECKLIST DE CONFORMIDADE
    # ══════════════════════════════════════════════════════════════════
    if checklist and isinstance(checklist, dict):
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        p_ck_title = doc.add_paragraph()
        _add_run(p_ck_title, "CHECKLIST DE CONFORMIDADE", bold=True, size=10, color=NAVY)
        p_ck_title.paragraph_format.space_after = Pt(4)

        for item, status in checklist.items():
            p_item = doc.add_paragraph()
            p_item.paragraph_format.space_after = Pt(1)
            p_item.paragraph_format.left_indent = Cm(0.5)
            icon = "✓" if status else "✗"
            icon_color = GREEN if status else RED
            _add_run(p_item, f"  {icon}  ", bold=True, size=10, color=icon_color)
            label = item.replace("_", " ").title()
            _add_run(p_item, label, size=10, color=BODY)

    # ══════════════════════════════════════════════════════════════════
    # ASSINATURA
    # ══════════════════════════════════════════════════════════════════
    doc.add_paragraph().paragraph_format.space_after = Pt(30)

    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_line, "___________________________________", size=11, color=LIGHT_GRAY)

    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_after = Pt(0)
    _add_run(p_name, medico_nome or "[Nome do Médico Responsável]", bold=True, size=11, color=DARK)

    p_crm = doc.add_paragraph()
    p_crm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_crm.paragraph_format.space_after = Pt(0)
    crm_line = medico_crm or "CRM: __________"
    if medico_rqe:
        crm_line = f"{crm_line} · RQE {medico_rqe}"
    _add_run(p_crm, crm_line, size=9, color=GRAY)

    # ── Watermark (rascunho) ──
    if not aprovado:
        doc.add_paragraph().paragraph_format.space_after = Pt(10)
        p_wm = doc.add_paragraph()
        p_wm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p_wm, "[ RASCUNHO — PENDENTE DE APROVAÇÃO ]", bold=True, size=12, color=LIGHT_GRAY)

    # ── Footer note ──
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_footer, f"Documento gerado em {datetime.now().strftime('%d/%m/%Y')} • Protocolo {protocolo} • Informações médicas confidenciais", size=8, color=LIGHT_GRAY)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
